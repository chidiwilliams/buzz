"""
Export to DOCX plugin for Buzz.

After a transcription completes, writes the transcript to a Microsoft Word
``.docx`` file next to the source media (or in a configured folder). Optionally
includes per-segment timestamps. Otherwise the text is grouped into paragraphs
using the same gap rule as Buzz's plain-text export.

Requires the ``python-docx`` package, declared as a pip dependency and installed
into the user cache on first load. The import is deferred to the hook so it does
not affect application startup.
"""

import logging
import os

from buzz.plugins.base import (
    BuzzPlugin,
    ConfigField,
    ConfigFieldType,
    PluginContext,
    PluginMetadata,
    plugin_gettext,
)

logger = logging.getLogger(__name__)

_ = plugin_gettext(__file__)

# Same default gap (ms) used by the TXT exporter to split paragraphs.
PARAGRAPH_SPLIT_TIME = 2000


class ExportDocxPlugin(BuzzPlugin):
    metadata = PluginMetadata(
        id="export_docx",
        name=_("Export to DOCX"),
        description=_(
            "Export the transcript to a Microsoft Word (.docx) file after "
            "transcription, optionally including timestamps."
        ),
        version="1.0.0",
        pip_dependencies=["python-docx"],
        config_fields=[
            ConfigField(
                key="output_folder",
                label=_("Output folder"),
                type=ConfigFieldType.TEXT,
                default="",
                description=_("Leave empty to save next to the source file."),
            ),
            ConfigField(
                key="include_timestamps",
                label=_("Include timestamps"),
                type=ConfigFieldType.BOOL,
                default=False,
            ),
        ],
    )

    def on_complete(self, transcription_id, task, segments, context: PluginContext):
        db_segments = context.transcription_service.get_transcription_segments(
            transcription_id=transcription_id
        )
        if not db_segments:
            context.log.info("Export to DOCX skipped: no segments")
            return

        source_path = task.file_path or task.original_file_path or "transcript"
        stem = os.path.splitext(os.path.basename(source_path))[0] or "transcript"

        folder = (context.config.get("output_folder") or "").strip()
        if not folder:
            folder = os.path.dirname(source_path) or os.getcwd()

        try:
            os.makedirs(folder, exist_ok=True)
            out_path = os.path.join(folder, f"{stem}.docx")
            self._write_docx(
                out_path,
                stem,
                db_segments,
                _coerce_bool(context.config.get("include_timestamps", False)),
            )
            context.log.info("Export to DOCX written to %s", out_path)
        except Exception as exc:
            context.log.error("Export to DOCX failed: %s", exc)

    def _write_docx(self, out_path, title, segments, include_timestamps):
        from docx import Document
        from buzz.transcriber.file_transcriber import to_timestamp

        document = Document()
        document.add_heading(title, level=1)

        if include_timestamps:
            for segment in segments:
                text = segment.text.strip()
                if not text:
                    continue
                stamp = (
                    f"[{to_timestamp(segment.start_time)} --> "
                    f"{to_timestamp(segment.end_time)}]"
                )
                paragraph = document.add_paragraph()
                run = paragraph.add_run(stamp + " ")
                run.bold = True
                paragraph.add_run(text)
        else:
            # Group into paragraphs on long gaps, mirroring the TXT exporter.
            current = []
            previous_end = None
            for segment in segments:
                if (
                    previous_end is not None
                    and (segment.start_time - previous_end) >= PARAGRAPH_SPLIT_TIME
                    and current
                ):
                    document.add_paragraph(" ".join(current))
                    current = []
                text = segment.text.strip()
                if text:
                    current.append(text)
                previous_end = segment.end_time
            if current:
                document.add_paragraph(" ".join(current))

        document.save(out_path)


def _coerce_bool(value) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.lower() in ("true", "1", "yes", "on")
    if isinstance(value, int):
        return value != 0
    return bool(value)
